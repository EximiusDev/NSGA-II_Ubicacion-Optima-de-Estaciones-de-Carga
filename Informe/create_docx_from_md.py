# -*- coding: utf-8 -*-
"""
Simple Markdown -> DOCX converter for the provided Informe text.
Supports:
- Headings: '## ' -> level 2, '### ' -> level 3
- Bullet lists starting with '- '
- Horizontal rules '---' (ignored, adds a blank line)
- Regular paragraphs

Saves: informe_nsga2.docx in the same folder.

Run: python create_docx_from_md.py

This script expects `python-docx` to be installed (`pip install python-docx`).
"""
from docx import Document
from docx.shared import Pt
import os

md_text = r"""
## 4. Experimentos y Resultados

### 4.1 Datos y preprocesamiento

Los datos se cargan desde `dataset_final.gpkg` con `geopandas.read_file`. Se garantiza un CRS métrico (`EPSG:3857`) para cálculos en metros. Las variables de prioridad (`flujo_total` y `esquinas_en_radio`) se normalizan dividiendo por su máximo observado y se combinan con pesos prefijados. Cuando las geometrías no son puntos, se usan centroides para representar la posición del nodo.

### 4.2 Configuración experimental

Parámetros representativos usados en los experimentos de referencia:

- `d_max` = 10,000 m (umbral de cobertura),  
- `pop_size` = 200,  
- `n_gen` = 500,  
- `factor_estaciones` = 50 (peso relativo en f1),  
- `max_ratio` = 0.15 (proporción máxima de nodos con estaciones permitida),  
- semilla (`seed`) = 23 para reproducibilidad.

Estas elecciones permiten equilibrar capacidad computacional y exploración del espacio de soluciones.

### 4.3 Principales resultados

- Frontera de Pareto: la ejecución de NSGA-II produjo una frontera con soluciones que trade-offeaban entre pocas estaciones y mejor cobertura. Las soluciones con baja cantidad de estaciones muestran un aumento notable en la distancia máxima cubierta, mientras que soluciones con más estaciones reducen fuertemente la distancia máxima.

- Prioridad: entre las soluciones con similar cobertura, algunas se distinguen por seleccionar nodos con mayor prioridad (mayor `flujo_total` y número de esquinas), lo que es deseable desde el punto de vista del servicio.

- Métricas de validación: la rutina `validar_soluciones` calcula por solución la distancia máxima real, cobertura promedio, % nodos cubiertos y eficiencia por estación. Estas métricas facilitan ordenar soluciones por criterios operativos.

(Nota: los resultados numéricos concretos están guardados en `resultado_nsga2_final.npz`. Para reproducir y visualizar las gráficas interactivas, ejecutar la celda de visualización en `resolucion_final_modificable.ipynb`. Si VS Code muestra advertencias de renderers para `plotly`, se recomienda usar `pio.renderers.default='browser'` o abrir el resultado en navegador.)

### 4.4 Interpretación

- Existe una curva de compromiso estable entre coste y cobertura. La elección operativa depende de recursos disponibles y objetivos de servicio.  
- Las soluciones que maximizan prioridad tienden a concentrar estaciones en nodos con alto flujo, lo que mejora la utilidad por estación pero puede dejar áreas periféricas con peor cobertura.  
- Sensibilidad: la frontera se desplaza notablemente al variar `d_max` y el factor de penalización; por tanto, se recomienda realizar análisis de sensibilidad antes de adoptar un diseño final.

---

## 5. Discusión y Limitaciones

- Distancia euclídea: la implementación usa distancia euclídea en un CRS métrico. Esto es una aproximación que puede ser insuficiente en redes viales complejas. Incorporar distancias por red (ruta real) mejora realismo.

- Costes y restricciones: no se modelaron costes reales de instalación, capacidad eléctrica ni restricciones urbanas. Estos factores son críticos para la implementación práctica.

- Escalabilidad: el cálculo de distancias puede ser costoso para grandes instancias. Se recomienda `cKDTree` para consultas de distancia mínima y paralelización o evaluación solo de soluciones promisorias.

- Datos: la calidad de `dataset_final.gpkg` (atributos y geometrías) condiciona fuertemente la validez de las recomendaciones. Se aconseja validar y enriquecer datos con tráfico, demanda temporal y límites regulatorios.

---

## 6. Conclusiones y recomendaciones (≈150–200 palabras)

Se desarrolló e implementó un marco basado en NSGA-II para ubicar estaciones de recarga de vehículos eléctricos sobre un conjunto de nodos geoespaciales. La formulación multiobjetivo permite explorar el conjunto de soluciones no dominadas, mostrando explícitamente los compromisos entre número de estaciones, cobertura y prioridad de servicio. Los resultados obtenidos, almacenados en `resultado_nsga2_final.npz` y producidos por `resolucion_final_modificable.ipynb`, evidencian que es posible reducir considerablemente la distancia máxima de cobertura aumentando moderadamente el número de estaciones, o priorizar nodos de alto flujo manteniendo un coste controlado. Sin embargo, para aplicar estas soluciones en entornos reales es necesario incorporar costes económicos, restricciones urbanas y distancias basadas en la red vial. Recomendamos usar el presente marco como herramienta de apoyo a la decisión, complementando la salida de NSGA-II con análisis de sensibilidad y validación con expertos locales. Futuras extensiones incluyen modelado de demanda temporal, integración de costes y optimización con rutas reales.

---

## Agradecimientos

Agradecemos a las autoridades y colegas que compartieron ideas y revisiones durante el desarrollo. El proyecto recoge contribuciones de recursos y datos incluidos en la carpeta del repositorio.

---

## Referencias (selección)

- Lam, A. Y. S., Leung, Y.-W., & Chu, X. Electric vehicle charging station placement: Formulation, complexity, and solutions.  
- Liu, Z.-f., Zhang, W., Ji, X., & Li, K. Optimal planning of charging station for electric vehicle based on particle swarm optimization.  
- Awasthi, A., Venkitusamy, K., Padmanaban, S., Selvamuthukumaran, R., Blaabjerg, F., & Singh, A. K. Optimal planning of electric vehicle charging station at the distribution system using hybrid optimization algorithm.  
- Bai, X., Chin, K.-S., & Zhou, Z. A bi-objective model for location planning of electric vehicle charging stations with GPS trajectory data.  
- Okabe, A., Sugihara, K., & Saeki, E. Kernel Density Estimation of Traffic Accidents in a Network Space.

---

## Apéndice A — Archivos relevantes en el repositorio

- `dataset_final.gpkg` — datos espaciales con nodos y atributos.  
- `resolucion_final_modificable.ipynb` — notebook con la implementación y experimentos.  
- `resultado_nsga2_final.npz` — resultados guardados (`X` y `F`).  
- `informe_nsga2.md`, `informe_nsga2.txt`, `informe_nsga2.rtf` — archivos auxiliares generados en el proyecto.

---

"""

def simple_md_to_docx(md, out_path):
    doc = Document()
    lines = md.splitlines()
    i = 0
    in_list = False
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.lstrip()
        if stripped == '':
            # blank line -> paragraph break
            in_list = False
            i += 1
            continue
        if stripped.startswith('### '):
            text = stripped[4:].strip()
            doc.add_heading(text, level=3)
            in_list = False
        elif stripped.startswith('## '):
            text = stripped[3:].strip()
            doc.add_heading(text, level=2)
            in_list = False
        elif stripped == '---':
            doc.add_paragraph('')
            in_list = False
        elif stripped.startswith('- '):
            # collect consecutive list items
            while i < len(lines) and lines[i].lstrip().startswith('- '):
                item = lines[i].lstrip()[2:].strip()
                p = doc.add_paragraph(item, style='List Bullet')
                i += 1
            in_list = False
            continue
        else:
            # normal paragraph — allow inline code markers to remain
            doc.add_paragraph(stripped)
            in_list = False
        i += 1

    # set a default style font size for Normal
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    doc.save(out_path)


if __name__ == '__main__':
    out_file = os.path.join(os.path.dirname(__file__), 'informe_nsga2.docx')
    print('Generando:', out_file)
    simple_md_to_docx(md_text, out_file)
    print('Listo. Archivo creado.')
